import pandas as pd
import openpyxl
from docx import Document
import os

path = r'D:\Cursos\Power BI\Proyecto Fitness\Data+gimnasio.xlsx'
output = r'Output/gym_data.txt'
excel_file = pd.ExcelFile(path)
document = Document()
#print(excel_file.sheet_names)
#df = pd.read_excel(path)



#Guardado
"""
with open(r'Output/gym_description.txt','w',encoding='utf-8') as file:
 pd.set_option('display.max_columns', None)
 for sheet in excel_file.sheet_names:
    data_sheet = pd.read_excel(path,sheet_name=sheet)
    file.write(f'sheet:{sheet}\n{data_sheet.head(5)}\n\n')
    

"""

for sheet in excel_file.sheet_names:
    data_sheet = pd.read_excel(path,sheet_name=sheet).head(5)
    nro = data_sheet.columns.value_counts()
    total_columns = sum(nro)
    document.add_heading(str(sheet),level=3)
    table = document.add_table(rows=1,cols=total_columns)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    for index, column in enumerate(data_sheet.columns):  
     
     header_cells[index].text = column
     
    for _, row in data_sheet.iterrows():
     row_cells = table.add_row().cells
     for idx, value in enumerate(row):
        row_cells[idx].text = str(value)

     


    
    

    #document.add_paragraph(data_sheet.head(5).to_string())

    
"""
for sheet in excel_file.sheet_names:
    data_sheet = pd.read_excel(path,sheet_name=sheet)

    document.add_heading(str(sheet),level=3)
    for index, column in enumerate(data_sheet.columns):
     #print(f"Índice: {index}, Nombre: {column}")
     print(data_sheet[column].head(5))

    
    document.add_paragraph(data_sheet.head(5).to_string())

    """  

if os.path.exists(r'Output/gym_description.docx'):
    os.remove(r'Output/gym_description.docx')

document.save(r"Output/gym_description.docx")