import csv
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt
import os
import openpyxl
from sklearn.linear_model import LinearRegression 

#documentation: https://python-docx.readthedocs.io/en/latest/

document = Document()
title = document.add_heading("Manipulacion de texto csv o txt", level=1)
subtitle = document.add_heading("Abrir archivos con csv", level=2)
paragraph = document.add_paragraph("")


source = r'D:\Cursos\Python\Manipulacion y limpieza\Resources\notebooks\datos.csv'
output_doc = r"Output/documento de practicas.docx"

if os.path.exists(output_doc):
    os.remove(output_doc)

"""Abrir archivos con csv"""

"""with open(source) as file:
    reader = csv.reader(file)
    for row in reader:"""
        

data = open(source)
reader = csv.reader(data)
for row in reader:
 paragraph = document.add_paragraph(' '.join(row))


#Abrir archivos con pandas
subtitle2 = document.add_heading("Abrir archivos con pandas", level=2)
data = pd.read_csv(source)
paragraph2 = document.add_paragraph(data.to_string())

comment = document.add_paragraph(r"para saltar tabulaciones agrega el parametro delimiter al pd.read asi: pd.read_csv(source, delimiter='\t')")

file_txt1 = pd.read_csv(r'Resources\notebooks\archivo1.txt',skiprows=2)


document.add_paragraph("Para archivos con comentarios con el signo #, se debe saltar las filas con el parametro skiprows = n")
document.add_paragraph(file_txt1.to_string())

document.add_paragraph("Dependiendo el caso agregar el parametro: engine='python'")


#Datos excel
source_excel = r'Resources\notebooks\datos_ejemplo.xlsx'

document.add_heading("Manipular archivos excel", level=2)
document.add_paragraph("para manipular archivos excel importar openpyxl")
document.add_paragraph("Si quieres acceder a todas las hojas de un documento excel, se puede hacer utilizando pd.ExcelFile")


xls = pd.ExcelFile(source_excel)

document.add_paragraph(' '.join(xls.sheet_names))

document.add_paragraph("si quieres imprimir todo el contenido de cada hoja se utiliza un bucle")


document.add_paragraph("Este es el resultado:")

for sheet_name in xls.sheet_names:
    data_excel = pd.read_excel(source_excel,sheet_name= sheet_name)
    document.add_paragraph(f'Contenido de la hoja {sheet_name}:')
    document.add_paragraph(data_excel.to_string())

document.add_paragraph("Para leer un grupo de columnas en especifico se utiliza la propiedad usecols, donde se especifica de que columna a que columna, ej: usecols='A:B'")

df_excel = pd.read_excel(source_excel)
document.add_paragraph("Para guardar un dataframe en un excel se convierte los datos en un dataframe y se utiliza la propiedad .to_excel, donde la propiendad index se deja en false si no se quiere agregar los numeros a las filas")

if os.path.exists('Output File/documento excel.xlsx'):
    os.remove('Output file/documento excel.xlsx')
    

df_excel.to_excel('Output/documento excel.xlsx',index=True)

#Descripcion de datos

document.add_heading("Descripcion preliminar de los datos", level=2)

data_txt = pd.read_csv(r'Resources\notebooks\datos.txt', delimiter='|')

document.add_paragraph('Para tener una descripcion de los datos utlizar el metodo describe de pandas ')
document.add_paragraph(data_txt.describe().to_string())

document.add_paragraph('Para tener un resumen de la tabla utilizar el metodo info:')

document.add_paragraph(data_txt.info())
document.add_paragraph('Si quiero saber los tipos utilizo la propiedad dtypes: ')
document.add_paragraph(data_txt.dtypes.to_string())

document.add_paragraph(r'Si quero obtener los valores unicos de una columna utilizo el metodo unique, especificando la columna, ejemplo: df["nombre"].unique():')

document.add_paragraph(str(data_txt['nombre'].unique()))

document.add_paragraph("Si quiero contarlos utilizo el metodo value_counts: ")
document.add_paragraph(data_txt['nombre'].value_counts().to_string())

#Datos faltantes
document.add_heading("Manejo de datos faltantes", level=2)
data_faltante = pd.read_csv(r'Resources/notebooks/datos_con_faltantes.csv')

document.add_paragraph("Para contar la cantidad de datos faltantes por columna utilizar los metodos .isnull().sum():")
document.add_paragraph(data_faltante.isnull().sum().to_string())

document.add_paragraph("Para eliminar los datos faltantes se utiliza el metodo .dropna()")

data_drop = data_faltante.dropna()

data_columns = data_faltante.dropna(axis= 1)
document.add_paragraph(data_drop.to_string())
document.add_paragraph("Si quiero eliminar las columnas con  datos faltantes se utiliza el parametro axis=1 en el metodo .dropna()")
document.add_paragraph(data_columns.to_string())

document.add_paragraph("Se puede rellenar los datos faltantes utilizando el metodo fillna() y especificar que deseas rellenar")

data_fill = data_faltante.copy()

data_faltante['Nombre'] = data_faltante['Nombre'].fillna("Desconocido")

document.add_paragraph(data_faltante.to_string())

document.add_paragraph("Para rellenar datos numericos como el salario se puede rellenar con la media")

data_faltante["Salario"] = data_faltante["Salario"].fillna(data_faltante["Salario"].mean())
document.add_paragraph(data_faltante.to_string())

document.add_paragraph("Para  datos numericos como la edad, es mas util  la mediana:")
data_faltante["Edad"] = data_faltante["Edad"].fillna(data_faltante["Edad"].median())
document.add_paragraph(data_faltante.to_string())

document.add_paragraph("Para rellenar datos como la cuidad, seria util el uso de la moda")
data_faltante["Ciudad"] = data_faltante["Ciudad"].fillna(data_faltante["Ciudad"].mode()[0])
document.add_paragraph(data_faltante.to_string())


document.add_paragraph('Para especificar los datos en donde se deban eliminar las columnas se utiliza la propiedad subset, ejemplo de subset= [Nombre]:')
document.add_paragraph(data_fill.dropna(subset=['Nombre'],inplace= False).to_string())

#Reestructurar datos
document.add_heading("Reestructurar datos",level=2)

document.add_paragraph("Para segmentar y clasificar datos continuos en grupos o intevalos se utiliza pd.cut():")
data_faltante["Rango_Edad"] = pd.cut(data_faltante["Edad"], bins=[20,25,30,35,40],labels=['20-25','25-30','30-35','35-40'],include_lowest= True)
document.add_paragraph(data_faltante.to_string())


document.add_paragraph("Para obtener el salario promedio por edad utilizo groupby() ")
data_agrupada = data_faltante.groupby('Rango_Edad',observed= True)['Salario'].mean()
document.add_paragraph(data_agrupada.to_string())

document.save("Output/documento de practicas.docx")

