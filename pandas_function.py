import pandas as pd
from docx import Document
from convert_to_table import to_table
import os



def pandas_manipulation():    
    document = Document()
    #Abrir archivos con pandas
    open_files_with_pd(document)
    
    #Descripcion preeliminar
    data_description(document)
    
    #Datos faltantes
    data_missing(document)
    
    #Reestructurar datos
    restructure_data(document)
    
    #Manejo de duplicados y Manejo de columnas 
    duplicate_and_columns_handling(document)

    #Concatenacion
    concatenation_handling(document)
    
    #Formato de largo a ancho e inverso
    format_lenght_width(document)

    #Separar columnas
    sepate_columns(document)

    #Conversion de datos categoricos
    categorical_data_conversion(document)

    #Variables dummy
    dummy_variables(document)

    #Funciones de agregación
    aggregation_function(document)

    #Funciones personalizadas
    custom_functions(document)
    
    #Funciones de grupo
    grouping_functions(document)

    
    #Exportación de datos
    export_data(document)

    #Manejo de json
    handling_json(document)

    #Manejo de xml y html
    handling_xml_html(document)

    #Uso de sql
    use_sql(document)

    #uso de orm
    use_orm(document)

    #Consumo de apis
    
    #Manejo de tiempo
    handling_time(document)
    
 
    document.save(r"Output/documento de practicas.docx")

    


def open_files_with_pd(document):
    source = r'D:\Cursos\Python\Manipulacion y limpieza\Resources\notebooks\datos.csv'
    
    document.add_heading(r"Abrir archivos con pandas", level=2)
    data = pd.read_csv(source)
    #crear tabla
    to_table(data, document)

    
    document.add_paragraph(r"""para saltar tabulaciones agrega el parametro delimiter al pd.read asi:
               pd.read_csv(source, delimiter='\t')""")
    
    file_txt1 = pd.read_csv(r'Resources\notebooks\archivo1.txt',skiprows=2)


    document.add_paragraph("""Para archivos con comentarios con el signo #, 
              se debe saltar las filas con el parametro skiprows = n""")
    
    to_table(file_txt1,document)

    document.add_paragraph("Dependiendo el caso agregar el parametro: engine='python'")

def data_description(document):
     
     document.add_heading("Descripcion preliminar de los datos", level=2)

     data_txt = pd.read_csv(r'Resources\notebooks\datos.txt', delimiter='|')

     document.add_paragraph('Para tener una descripcion de los datos utlizar el metodo describe de pandas ')
     document.add_paragraph(data_txt.describe().to_string())


     document.add_paragraph('Si quiero saber los tipos utilizo la propiedad dtypes: ')
     document.add_paragraph(data_txt.dtypes.to_string())

     document.add_paragraph(r"""Si quero obtener los valores unicos de una columna utilizo el metodo unique, e
                specificando la columna, ejemplo: df["nombre"].unique():""")

     document.add_paragraph(str(data_txt['nombre'].unique()))

     document.add_paragraph("Si quiero contarlos utilizo el metodo value_counts: ")
     document.add_paragraph(data_txt['nombre'].value_counts().to_string())

def data_missing(document):
    document.add_heading("Manejo de datos faltantes", level=2)
    data_faltante = pd.read_csv(r'Resources/notebooks/datos_con_faltantes.csv')

    document.add_paragraph("""Para contar la cantidad de datos faltantes por columna 
              utilizar los metodos .isnull().sum():""")
    document.add_paragraph(data_faltante.isnull().sum().to_string())

    document.add_paragraph("Para eliminar los datos faltantes se utiliza el metodo .dropna()")

    data_drop = data_faltante.dropna()

    #Crear tabla
    to_table(data_drop,document)

    data_columns = data_faltante.dropna(axis= 1)
    
    document.add_paragraph("""Si quiero eliminar las columnas con  datos faltantes se utiliza el parametro axis=1
               en el metodo .dropna()""")
    
    to_table(data_columns,document)
    

    document.add_paragraph("""Se puede rellenar los datos faltantes utilizando el metodo fillna()
               y especificar que deseas rellenar""")

    data_fill = data_faltante.copy()

    data_faltante['Nombre'] = data_faltante['Nombre'].fillna("Desconocido")

    to_table(data_faltante,document)
    

    document.add_paragraph("Para rellenar datos numericos como el salario se puede rellenar con la media")

    data_faltante["Salario"] = data_faltante["Salario"].fillna(data_faltante["Salario"].mean())
    to_table(data_faltante,document)

    document.add_paragraph("Para  datos numericos como la edad, es mas util  la mediana:")
    data_faltante["Edad"] = data_faltante["Edad"].fillna(data_faltante["Edad"].median())
    to_table(data_faltante,document)

    document.add_paragraph("Para rellenar datos como la cuidad, seria util el uso de la moda")
    data_faltante["Ciudad"] = data_faltante["Ciudad"].fillna(data_faltante["Ciudad"].mode()[0])
    to_table(data_faltante,document)


    document.add_paragraph("""Para especificar los datos en donde se deban eliminar las columnas
                            se utiliza la propiedad subset, ejemplo de subset= [Nombre]:""")
    to_table(data_fill.dropna(subset=['Nombre'],inplace= False), document)

def restructure_data(document):

    document.add_heading("Reestructurar datos", level=2)
    data_faltante = pd.read_csv(r'Resources/notebooks/datos_con_faltantes.csv')
    document.add_paragraph("Para segmentar y clasificar datos continuos en grupos o intevalos se utiliza pd.cut():")
    data_faltante["Rango_Edad"] = pd.cut(data_faltante["Edad"], bins=[20,25,30,35,40],labels=['20-25','25-30','30-35','35-40'],include_lowest= True)
    to_table(data_faltante, document)


    document.add_paragraph("Para obtener el salario promedio por edad utilizo groupby() ")
    data_agrupada = data_faltante.groupby('Rango_Edad',observed= True)['Salario'].mean()
    document.add_paragraph(data_agrupada.to_string())

def duplicate_and_columns_handling(document):
 duplicate = pd.read_csv(r'Resources\notebooks\data_duplicada.csv')

 document.add_heading("Manejo de duplicados", level=2)
 document.add_paragraph("""Al usar el metodo .duplicated() me indica por medio de booleans que filas son duplicadas, sin necesidad de que estos sean consecutivos""")

 document.add_paragraph(duplicate.duplicated().to_string())

 document.add_paragraph("""Se puede usar el parametro subset para especificar las columnas y saber si contiene duplicados""")
    
 document.add_paragraph(duplicate.duplicated(subset='Nombre').to_string())

 document.add_paragraph("""Obteniendo las filas de los duplicados con: duplicate[duplicate.duplicated()""")

 duplicated_rows = duplicate[duplicate.duplicated()]
 to_table(duplicated_rows, document)

 document.add_paragraph("""Para eliminar los duplicados se utiliza el metodo .drop_duplicates()""")
 duplicate_delete = duplicate.drop_duplicates()
 to_table(duplicate_delete,document)

 document.add_paragraph("Para crear una columna que indique los valores duplicados se hace de este metodo: duplicate['Es duplicado?'] = duplicate.duplicated() ")
 duplicate['Es duplicado?'] = duplicate.duplicated()
 to_table(duplicate,document)

 document.add_paragraph("Usando .map puedo cambiar los valores de la columna 'Es duplicado' de un booleano a Si 0 No")
 duplicate["Es duplicado?"] = duplicate["Es duplicado?"].map({True: 'Si', False: 'No'})
 to_table(duplicate,document)
 """ duplicate_summarized = duplicate.groupby('Nombre').agg({
        'Edad': 'first',
        'Salario': 'mean',
        'Fecha_Ingreso': 'first'
    }).reset_index()"""
  #Manejo de columnas
 document.add_paragraph("Para reordenar las columnas de un dataframe se hace de esta forma:duplicate[['ID','Nombre','Edad', 'Salario','Es duplicado?','Fecha_Ingreso']]")
 to_table(duplicate[['ID','Nombre','Edad', 'Salario','Es duplicado?','Fecha_Ingreso']],document)
    
 document.add_paragraph("Para elegir columnas especificas de un dataframe se puede usa la propiedad .loc de esta manera: duplicate.loc[:,['ID','Nombre']] ")
 duplicate_loc = duplicate.loc[:,["ID",'Nombre']]
 to_table(duplicate_loc,document)
 document.add_paragraph("El primer parametro de la propiedad .loc se utiliza para especificar el rango de las filas que se selecionaran, ejemplo de filas del 1 al 9 / 1:9")
 to_table(duplicate_loc[1:9],document)

 document.add_paragraph("Para eliminar una columna, ejemplo la duplicado, se utiliza el metodo ..drop(columns=['Es duplicado?'])")
 duplicate = duplicate.drop(columns=['Es duplicado?'])
 to_table(duplicate,document)
 document.add_paragraph("Si quiero filtra los salarios superiores a 50,000 se hace de la siguiente manera: duplicate.loc[duplicate['Salario' ]> 50000]")
    
 to_table(duplicate.loc[duplicate['Salario' ]> 50000],document)

 document.add_paragraph("Para agregar una columna nueva a un dataframe se hace de esta manera:duplicate['Posición'] agregando los valores que tendra la columna: ")
 duplicate['Posición'] = pd.cut(duplicate["Salario"],bins=[40000,60000,65000,75000],labels=['junior','mid','Senior'])
 to_table(duplicate,document)

 document.add_paragraph("Esta nueva columna se calcula cuanto se le descuenta de afp + ars: ")
 duplicate['AFP + ARS'] = duplicate["Salario"] * (5.91 /100)
 duplicate['Salario_Neto'] = duplicate['Salario'] * 0.9409
 to_table(duplicate,document)

def concatenation_handling(document):
    document.add_heading("Concatenacion y combinacion", level=2)
    
    data_lima = {
    'Producto': ['A', 'B'],
    'Ventas': [250, 150],
    'Ciudad': ['Lima', 'Lima']
    }

    df_lima = pd.DataFrame(data_lima)
    
    
    data_bogota = {
    'Producto': ['A', 'C'],
    'Ventas': [100, 200],
    'Ciudad': ['Bogotá', 'Bogotá']
}
    df_bogota = pd.DataFrame(data_bogota)
    document.add_paragraph("Para concatenar dos o mas dataframes se utiliza el metodo concat de esta manera: pd.concat([df_lima, df_bogota]), este es el resultado: ")
    df_concat = pd.concat([df_lima, df_bogota])
    to_table(df_concat, document)
    data_inventario = {
    'Producto': ['A', 'B'],
    'Inventario': [30, 45]
    }

    df_inventario = pd.DataFrame(data_inventario)
    document.add_paragraph("Al concatenar datframes es recomendable reiniciar indices usando la propiedad: reset_index(drop=True)  ")
    df_concat = pd.concat([df_concat.reset_index(drop=True), df_inventario],axis=1)
    to_table(df_concat,document)

def format_lenght_width(document):
    document.add_heading("Formatear tablas de largo a ancho", level=2)
    data_largo = {
    'Producto': ['A', 'A', 'B', 'B', 'C', 'C'],
    'Mes': ['Enero', 'Febrero', 'Enero', 'Febrero', 'Enero', 'Febrero'],
    'Ventas': [100, 150, 200, 250, 300, 350]
    }

    df = pd.DataFrame(data_largo)
    document.add_paragraph("Para convertir un dataframe de largo a ancho se utiliza el metodo pivot de un dataframe, que te pide tres parametros, index, columns y values")
    document.add_paragraph("Antes de utilizar el metodo pivot:")
    to_table(df,document)
    document.add_paragraph("Despues de utilizar el metodo pivot:")
    document.add_paragraph(df.pivot(index='Producto', columns='Mes', values='Ventas').to_string())
    document.add_heading("Formatear tablas de ancho a largo", level=2)
    df_width = df.pivot(index='Producto', columns='Mes', values='Ventas')
    document.add_paragraph("Para formatear dataframes, para pasarlos de ancho a largo se utiliza el metodo melt, el cual te pide un dataframe y que llenes los parametros: id_vars,value_vars,var_name, value_name, el resultado es este:")
    df_new_lenght = pd.melt(df_width.reset_index(),id_vars=['Producto'],value_vars=['Enero', 'Febrero'],var_name='Mes', value_name='Ventas')
    to_table(df_new_lenght,document)

def sepate_columns(document):

    document.add_heading("Separar columnas", level=2)
    data = {
    'Nombre_Completo': ['Juan Perez', 'Maria Gomez', 'Luis Martinez']}
    data_fechas = {
    'Fecha': ['01-01-2024', '15-02-2024', '30-03-2024']
}

    
    df = pd.DataFrame(data)
    df_fechas = pd.DataFrame(data_fechas)
    document.add_paragraph("Para separar los datoscde una tabla de un dataframe en columnas se utiliza el metodo split de la propiedad str: .str.split, donde en el primer parametro se expecifica que elemento se debe separar del texto. ")
    document.add_paragraph("Antes de separar:")
    to_table(df,document)
    document.add_paragraph("\n")
    to_table(df_fechas,document)
    
    document.add_paragraph("Despues de separar:")
    df[['Nombre', 'Apellido']] = df['Nombre_Completo'].str.split(' ', expand=True)
    df_fechas[['dia','mes','año']] = df_fechas['Fecha'].str.split('-', expand=True)
    
    to_table(df,document)
    document.add_paragraph("\n")
    to_table(df_fechas,document)
    document.add_paragraph("Para crear una fecha completa con las / se usa el metodo agg, y el metodo join de esta manera: .agg('/'.join,axis=1) ")
    df_fechas['Fecha_completa'] = df_fechas[['dia','mes','año']].agg('/'.join,axis=1)
    document.add_paragraph("Resultado:")
    to_table(df_fechas,document)  

def categorical_data_conversion(document):
    document.add_heading("Conversion de datos categoricos", level=2)
    data = {
    'producto': ['Manzana', 'Banana', 'Cereza', 'Durazno', 'Pera'],
    'categoria': ['Fruta', 'Fruta', 'Fruta', 'Fruta', 'Fruta'],
    'calidad': ['Alta', 'Media', 'Baja', 'Alta', 'Media'],  # Columna categórica
    'ventas': [50, 30, 70, 85, 40]
}

    document.add_paragraph("Para la conversion de datos categoricos se puede utilizar el metodo de codificacion ordinal que consiste en darle un valor a las variables categoricas para poder darles un orden  lógico o jerárquico  ")
    df = pd.DataFrame(data)
    document.add_paragraph("Datatframe si en el orden categorico:")
    to_table(df,document)
    category_map = {'Baja':1, 'Media':2, 'Alta':3}
    df['Categoria ordinal'] = df['calidad'].map(category_map)
    document.add_paragraph("Despues:")
    to_table(df,document)

    document.add_paragraph("La codificacion one hot permite convertir cada columna categoria a binaria:")

    df_one_hot = pd.get_dummies(df,columns=['calidad'])
    to_table(df_one_hot,document)

def dummy_variables(document):
    document.add_heading("Variables dummy", level=2)
    data = {
    'vehiculo': ['Auto', 'Camioneta', 'Moto', 'Camion', 'Auto'],
    'color': ['Rojo', 'Azul', 'Negro', 'Blanco', 'Rojo'],
    'precio': [20000, 30000, 15000, 40000, 18000],
    'ventas': [150, 120, 130, 60, 180]
     }

    df = pd.DataFrame(data)
    document.add_paragraph("Las variables dummy son variables categoricas que toman valores binarios, usando el metodo get_dummies se crean columnas binarias con estos valores")
    document.add_paragraph("Antes de la conversión:")
    to_table(df,document)
    dummy_variables = pd.get_dummies(df,columns=['vehiculo', 'color'])
    document.add_paragraph("Despues:")
    to_table(dummy_variables,document)
    document.add_paragraph("Usando el parametro drop_first, se elimina la multicolinealidad que consiste en eliminar las columnas que son redundates")
    dummy_variables_drop = pd.get_dummies(df, columns=['vehiculo','color'],drop_first=True)
    to_table(dummy_variables_drop,document)
    document.add_paragraph("Si quiero que los valores de las columnas dummy se muestren con texto descriptivos, se puede usar el metodo map, pero usando una expresion lambda en vez de booleanos:")
    rename_dummy_values = dummy_variables_drop[['vehiculo_Camion', 'vehiculo_Camioneta','vehiculo_Moto','color_Blanco', 'color_Negro','color_Rojo']].map(lambda x: 'Si' if x else 'No')
    to_table(rename_dummy_values,document)

def aggregation_function(document):
    document.add_heading("Funciones de agregación",level=2)
    
    data = {
    'producto': ['Manzana', 'Manzana', 'Banana', 'Banana', 'Cereza', 'Cereza', 'Manzana', 'Banana', 'Cereza'],
    'categoría': ['Fruta', 'Fruta', 'Fruta', 'Legumbre', 'Fruta', 'Fruta', 'Fruta', 'Legumbre', 'Fruta'],
    'ventas': [50, 30, 20, 15, 10, 25, 45, 55, 40],
    'precio_unitario': [0.5, 0.5, 0.3, 0.3, 0.8, 0.8, 0.5, 0.3, 0.8],
    'fecha': pd.date_range(start='2024-01-01', periods=9, freq='D')
    }
    document.add_paragraph("Para las funciones de agregacion usando el metodo groupby se debe especificar la columna categorica, los valores los cuales seran agrupados y el calculo que se realizara")
    df = pd.DataFrame(data)
    document.add_paragraph("Datos sin agrupar")
    to_table(df,document)
    document.add_paragraph("Datos de suma de ventas agrupados por producto, df.groupby('producto')['ventas'].sum():  ")
    product_sales = df.groupby('producto')['ventas'].sum()
    document.add_paragraph(product_sales.to_string())
    document.add_paragraph("Datos del promedio de precio por unidad agrupados por producto, df.groupby('producto')['precio_unitario'].mean() ")
    average_product_price = df.groupby('producto')['precio_unitario'].mean()
    document.add_paragraph(average_product_price.to_string())
    document.add_paragraph("Con el metodo agg se puede agregar multiples funciones, ejemplo:")
    sales_summary = df.groupby('producto')['ventas'].agg(['sum','min','max'])
    to_table(sales_summary,document)
    document.add_paragraph("Si quiero cambiar los datos agrupados convertidos en una serie a un dataframe se utiliza el metodo .reset_index() ")
    sales_summary = df.groupby('producto')['ventas'].agg(['sum','min','max']).reset_index()
    to_table(sales_summary,document)
    document.add_paragraph("Se puede agrupar varias columnas con groupby y crear varias funciones con el metodo agg: ")
    category_summary = df.groupby(['categoría','producto']).agg(
        total_sales = ('ventas','sum'),
        average_price = ('precio_unitario','mean')
    )
    to_table(category_summary,document)


def custom_functions(document):
    document.add_heading("Funciones personalizadas", level=2)

    data = {
    'producto': ['Manzana', 'Manzana', 'Banana', 'Banana', 'Cereza', 'Cereza', 'Manzana', 'Banana', 'Cereza'],
    'categoría': ['Fruta', 'Fruta', 'Fruta', 'Fruta', 'Fruta', 'Fruta', 'Fruta', 'Fruta', 'Fruta'],
    'ventas': [50, 30, 20, 15, 10, 25, 45, 55, 40],
    'precio_unitario': [0.5, 0.5, 0.3, 0.3, 0.8, 0.8, 0.5, 0.3, 0.8],
    'fecha': pd.date_range(start='2024-01-01', periods=9, freq='D')
}

    df = pd.DataFrame(data)

    document.add_paragraph("Usan el metodo apply, se puede integrar funciones personalizadas")
    document.add_paragraph("Dataframe original:")
    to_table(df,document)


    def total_renevue(row):
        return row['ventas'] * row['precio_unitario']
    
    df['Ingresos totales'] = df.apply(total_renevue,axis=1)
    document.add_paragraph("dataframe con los ingresos totales:")
    to_table(df,document)

    def classification_sales(ventas):
        classification = {
            'Alto': ventas > 40,
            'Medio': 20 <= ventas <= 40,
            'Bajo': ventas < 20
        }
        for categoria, cantidad in classification.items():
            if cantidad:
                return categoria

    document.add_paragraph("dataframe con la clasificacion de las ventas:")
    df['clasificaciones'] = df['ventas'].apply(classification_sales)
    to_table(df,document)
    
    def percentage_sales(ventas):
        total = ventas.sum()
        porcentaje = (ventas / total) * 100
        return porcentaje.astype(int)
    
    document.add_paragraph("dataframe con el porcentaje de las ventas:")
    df['Porcentaje'] = df.groupby('producto')['ventas'].transform(percentage_sales)
    to_table(df,document)

def grouping_functions(document):
    document.add_heading("Funciones de grupo", level=2)
    document.add_paragraph("Estas funciones por grupo esta hecho con los metodos groupby y agg")

    data = {
    'producto': ['Manzana', 'Manzana', 'Banana', 'Banana', 'Cereza', 'Cereza', 'Manzana', 'Banana', 'Cereza'],
    'categoría': ['Fruta', 'Fruta', 'Fruta', 'Fruta', 'Fruta', 'Fruta', 'Fruta', 'Fruta', 'Fruta'],
    'ventas': [50, 30, 20, 15, 10, 25, 45, 55, 40],
    'precio_unitario': [0.5, 0.5, 0.3, 0.3, 0.8, 0.8, 0.5, 0.3, 0.8],
    'fecha': pd.date_range(start='2024-01-01', periods=9, freq='D')
}

    df = pd.DataFrame(data)

    
    document.add_paragraph("Dataframe original:")
    document.add_paragraph(df.to_string())

    grouping_products_by_sales = df.groupby('producto')['ventas'].agg(['sum','mean']).reset_index()
    document.add_paragraph("Agrupación de productos por ventas:")
    document.add_paragraph(grouping_products_by_sales.to_string())

    def range_function(ventas):
        return ventas.max() - ventas.min()

    range_sales = df.groupby('producto')['ventas'].agg(range_function).reset_index()
    document.add_paragraph("Rango de ventas, (hecho con una función personalizada):")
    document.add_paragraph(range_sales.to_string())
    document.add_paragraph("Agrupación de ingreso total por producto:")
    df['ingreso_total'] = df['ventas'] * df['precio_unitario']
    
    grouping_reneuve = df.groupby(['producto','categoría']).agg({
        'ventas': 'sum',
        'ingreso_total': 'sum'
    }).reset_index()

    document.add_paragraph(grouping_reneuve.to_string())

def export_data(document):
    document.add_heading("Exportación de datos", level=2)
    document.add_paragraph("""Para exportar datos se debe primero convertir el conjunto de datos en un dataframe, despues se elige el formato para exportar, de los cuales estan:
    to_csv
    to_excel
    to_parquet
    to_html
    etc.
                           """)
    document.add_paragraph("Al elegir cualquiera de los formatos para exportar se debe ingresar el nombre que tendra el archivo y su extensión de esta manera: to_excel('Output/documento excel.xlsx',index=True) donde la propiedad index indica si deberia estar numerada cada fila del dataframe. ")
    document.add_paragraph("""Recomendaciones para el guardado de archivos:
    En archivos csv, el parametro sep es para separar los delimitadores.
    float_format='%.2f' Para expecificar que solo haya 2 decimales
    El parametro compression es para especificar el formato en el que se comprimirá el archivo
    el parametro compression permite los siguientes procesos compression={'method': 'zip', 'archive_name': 'datos_comprimido.xlsx'}
    el parametro encoding, permite codificar el archivo a multiples formatos: encoding='utf-8'
    Con el parametro na_rep puedes especificar con que rellenar los datos nulos o vacios, ej: na_rep='N/A'        """)

def handling_json(document):
    document.add_heading("Manejo de json", level=2)
    document.add_paragraph("Para pasar cadenas de json al metodo  pd.read_json, es necesario convertir esas cadenas en un formato StringIO de la siguiente forma: ")
    document.add_paragraph("""json_data = '''
[
    {"producto": "A", "ventas": 100},
    {"producto": "B", "ventas": 200},
    {"producto": "C", "ventas": 300}
]
'''

# Envolver la cadena JSON en un objeto StringIO
json_stream = StringIO(json_data)

Es necesario hacer estas importaciones: from io import StringIO """)
    

    json = pd.read_json(r'Resources\notebooks\datos.json',orient='records')
    document.add_paragraph("Diferentes orientaciones que se pueden utilizar a exportar en json:")
    document.add_paragraph("orient='split':")
    format_split = json.to_json(orient='split')
    document.add_paragraph(str(format_split))
    document.add_paragraph("orient='records' uso de la propiedad lines = True (unicamente utilizable con esta orientación):")
    document.add_paragraph(str(json.to_json(orient='records', lines= True)))
    document.add_paragraph("orient='index'")
    document.add_paragraph(str(json.to_json(orient='index')))
    document.add_paragraph("orient='columns:'")
    document.add_paragraph(str(json.to_json(orient='columns')))
    document.add_paragraph("orient='values':")
    document.add_paragraph(str(json.to_json(orient='values')))

def handling_xml_html(document):
    document.add_heading("Manejo de xml y html",level=2)
    document.add_paragraph("Para la lectura de archivos xml se usa el metodo read_xml, si no es un archivo, se usa StringIO de la libreria io al igual que los json  ")
    df = pd.read_xml(r'Resources\notebooks\datos.xml')
    document.add_paragraph(df.to_string())
    
    tables = pd.read_html(r'Resources\notebooks\datos.html')
    document.add_paragraph("Si un documento html posee varias tablas usando tables[n] donde n es el numero de tabla que se busca acceder:")
    document.add_paragraph(tables[0].to_string())

def use_sql(document):
    document.add_heading("Uso de sql", level=2)
    from sqlalchemy import create_engine
    data = {
    'nombre': ['Juan', 'Ana', 'Luis', 'María', 'Pedro'],
    'apellido': ['Pérez', 'López', 'González', 'Rodríguez', 'Martínez'],
    'edad': [30, 28, 40, 35, 45],
    'salario': [50000, 55000, 70000, 65000, 80000],
    'departamento': ['Ventas', 'Marketing', 'Finanzas', 'Ventas', 'Recursos Humanos']
}

    df = pd.DataFrame(data)
    document.add_paragraph("Para crear un base de datos de sqlite se usa el metodo create_engine de sqlalchemy, el parametro que exige es este: sqlite:///Output/nombre_de_la_basededatos.db, (Las primeras 3 diagonales son obligatorias cuando no hay un nombre de usuario, contraseña o demas parametros)")

   
    

    engine = create_engine('sqlite:///Output/empleados.db')
    df.to_sql('empleados',con=engine,if_exists='replace',index=False)

    document.add_paragraph("Con el metodo to_sql de un dataframe puedes hacer que se guarden los datos en una tabla de una base de datos.")
    document.add_paragraph("""Parametros:
    name: (obligatorio), Especifica el nombre de la tabla en la base de datos donde se guardará el DataFrame.
    con: (obligatorio), La conexión a la base de datos, creada previamente con SQLAlchemy (usando create_engine).
    if_exists: Especifica que hacer si la tabla ya existe.
    index: Pametro booleano para indicar si se debe guardar el indice en la base de datos""")

    document.add_paragraph("Con el metodo read_sql se puede pasar una query para realizar una operación en db, toma como parametro la conexión 'con'")
    response = pd.read_sql('select * from empleados', con = engine)
    
    #print(response)

def use_orm(document):
    from empleado import Empleado, session
    document.add_heading("Uso de orm", level=2)
    document.add_paragraph("La libreria de sqlalchemy posee multiples clases para mapear clases de python en tablas para una base de datos, entre ellas estan Column, que representaran las columnas de las tablas")
    document.add_paragraph("""De sqlalchemy.orm estan:
     declarative_base: Clase base para definir modelos (tablas).
     sessionmaker: Crea sesiones para interactuar con la base de datos.                      
                           """)
    document.add_paragraph("Al instanciar Base = declarative_base() Base se convierte en una clase la cual servira para definir las clases que seran mapeadas en tablas para las base de datos. ")
    document.add_paragraph("""Con Session se realiza lo siguiente: 
    Session = sessionmaker(bind=engine)
    session = Session()
                        
    sessionmaker: Genera una fábrica de sesiones ligadas al motor (base de datos).
    session: Es una instancia de sesión que se usa para interactuar con la base de datos.
""")
    document.add_paragraph("Con:  Base.metadata.create_all(engine) Crea las tablas en la base de datos basadas en los modelos definidos (si no existen).")
    document.add_paragraph("""Con: session.add(empleado) y session.commit() ocurre lo siguiente:
    session.add(empleado): Añade la instancia al contexto de la sesión.
    session.commit(): Confirma los cambios y guarda el registro en la base de datos.

        
""")

    #new_empleo = Empleado(nombre = 'Griselda', apellido='Matos', edad=22, salario=25000, departamento='RH')
    #session.add(new_empleo)
    session.commit()
    """
    empleados = session.query(Empleado).all()
    for empleado in empleados :
        
        print(empleado)
    """

def handling_time(document):
    document.add_heading("Manejo de fechas y tiempo",level=2)
    document.add_paragraph("""El metodo date_range de pandas permite generar una serie de fechas apartir de una fecha inicial, sus parametros son:
    start: Recibe un valor de fecha o de tipo string que sea una fecha, sirve como fecha inicial.
                           
    periods: Es un valor de tipo integer que especifica la cantidad de fechas o puntos de tiempo se deben generar en el rango
    freq: Especifica la frecuencia de los intervalos entre las fechas generadas. Por defecto es 'D' (diaria).
""")
    fechas = pd.date_range(start='2023-01-01', periods=10, freq='D')
    precios = [150, 152, 149, 153, 155, 158, 157, 160, 162, 161]
    serie_temporal = pd.Series(data=precios, index=fechas)
    document.add_paragraph("Una serie es una estructura de datos unidimensional. se pueden crear series con el metodo pd.Series de pandas, posee como parametros data que recibe la información de la serie e index, que pueden contener los idices de los datos.")
    document.add_paragraph(serie_temporal.to_string())
    document.add_paragraph("Con la propiedad .iloc[n] se puede obtener el elemento de una serie basado en su posición numerica, si se usa un slicing solo traera los datos de la primera posición especificada hasta la anterior a la ultima espcificada, ej: 1:3 donde se seleccionará los elementos en las posiciones 1 y 2 de la Serie. ")
    document.add_paragraph(str(serie_temporal.iloc[1]))
    document.add_paragraph(serie_temporal.iloc[1:3].to_string())
    document.add_paragraph("Se puede obtener un subconjunto de datos de las series usando slicing ya sea espicificando el indice o la posición.")
    document.add_paragraph("Por posicion [0:3]")
    document.add_paragraph(serie_temporal[0:3].to_string())
    document.add_paragraph("Por indice '2023-01-01':'2023-01-05'")
    document.add_paragraph(serie_temporal['2023-01-01':'2023-01-05'].to_string())
    
    document.add_paragraph("Con el metodo diff puedes obtener la diferencia de valores de los datos con respecto a la fecha anterior.")
    document.add_paragraph(serie_temporal.diff().to_string())
    
    document.add_paragraph("El metodo resample se utiliza para reasignar o agrupar datos de series temporales en diferentes intervalos de tiempo, acepta como parametro el tipo de frecuencia al cual se quiere cambiar ej: (dias,semanas, meses etc)")
    document.add_paragraph("Suma de series agrupadas en semana: ")
    serie_semanal = serie_temporal.resample('W').sum()
    document.add_paragraph(serie_semanal.to_string())

    fechas = ['2024-08-01', '01/08/2024', 'August 1, 2024', '20240801']
    df = pd.DataFrame({'fecha': fechas})
    document.add_paragraph("Para el formateo de fechas se usa el metodo  pd.to_datetime función que convierte una columna de tipo texto, enteros u otros formatos a un objeto de fecha y hora (datetime64)., con el parametro format se puede asignar un formato especifico a las fechas.")
    document.add_paragraph("Fechas del dataframe sin formatear")
    document.add_paragraph(df.to_string())
    
    df['fecha'] = pd.to_datetime(df['fecha'],format='mixed')
    document.add_paragraph("Fechas convertidas con to_datetime y formateadas con format='mixed' (Formato mixto)")
    document.add_paragraph(df.to_string())
    

    document.add_paragraph("con el metodo .dt.strftime('%d-%m-%Y') puedes formatear columnas en formato datetime en columnas string formateadas en una patron especifico")
    df['fecha_formateada'] = df['fecha'].dt.strftime('%d-%m-%Y')
    document.add_paragraph("Fecha formateada en '%d-%m-%Y' dia-mes-año: ")
    document.add_paragraph(df.to_string())
    

    document.add_paragraph("El argumento errors='coerce' en el método pd.to_datetime() se utiliza para manejar errores durante la conversión de datos a formato de fecha y hora (datetime). Si pandas encuentra un valor que no puede ser convertido a una fecha válida, en lugar de generar un error, asignará un valor especial llamado NaT ('Not a Time') ")
    df = df.drop(columns='fecha_formateada')
    document.add_paragraph("Para la extracción de partes de una fecha (como dia mes o año) se utiliza dt.day|dt.month|dt.year")
    df['dia'] = df['fecha'].dt.day
    df['mes'] = df['fecha'].dt.month
    df['año'] = df['fecha'].dt.year
    document.add_paragraph(df.to_string())
    

    fechas_horas = pd.to_datetime(['2024-08-01 14:30:00', '2023-12-15 09:45:30', '2025-07-20 22:10:15'])
    df_horas = pd.DataFrame({'fecha_hora': fechas_horas})

    
    document.add_paragraph("Para la extraccion del tiempo (hora, minuto, segundo) se utiliza dt.hour|dt.minute|dt.second")
    document.add_paragraph("Fechas con horas:")
    document.add_paragraph(df_horas.to_string())
    
    df_horas['hora'] = df_horas['fecha_hora'].dt.hour
    df_horas['minuto'] = df_horas['fecha_hora'].dt.minute
    df_horas['segundo'] = df_horas['fecha_hora'].dt.second
    document.add_paragraph(df_horas.to_string())

    df_horas['dia_de_la_semana'] = df_horas['fecha_hora'].dt.weekday
    df_horas['dia_del_año'] = df_horas['fecha_hora'].dt.day_of_year
    df_horas['dia_semana_nombre'] = df_horas['fecha_hora'].dt.day_name()
    df_horas['trimestre'] = df_horas['fecha_hora'].dt.quarter

    #print(df_horas[['fecha_hora','dia_de_la_semana','dia_del_año']])
    

    document.add_paragraph("El método dt.isocalendar() se utiliza para obtener información relacionada con el calendario ISO para una columna de fechas. El método devuelve un DataFrame con tres columnas: year/week/day:")
    document.add_paragraph(df_horas['fecha_hora'].dt.isocalendar().to_string())
   

    document.add_paragraph("Si una columna con valores datetime posee datos faltantes se pueden usar los metodos ffill() y bffill(). ffill rellena la información faltante con la fecha de la fila anterior mientras que bffill con la posterior")
    document.add_paragraph("Con pd.Timestamp puedes convertir una cadena en un objeto de tipo Timestamp (fecha y hora en pandas). Combinándolo con .fillna() puedes remplazar los valores faltantes con la fecha proporcionada ")
    

    fechas = ['2024-08-01', None, '2024-99-99', '2023-12-15']
    df = pd.DataFrame({'fecha': fechas})
    df['fecha'] = pd.to_datetime(df['fecha'], errors='coerce')
    document.add_paragraph("Fecha sin limpiar:")
    document.add_paragraph(df.to_string())

    
    document.add_paragraph("uso de drona():")
    df_cleaned = df.dropna(subset=['fecha'])
    document.add_paragraph(df_cleaned.to_string())
    document.add_paragraph("Uso de notna() devuelve un booleano que indica  si hay valores vacios o no")
    document.add_paragraph(df['fecha'].notna().to_string())

    document.add_paragraph("Se puede hacer calculos con las fecha como ejemplo restar:")
    df = pd.DataFrame({
    'fecha_inicio': pd.to_datetime(['2024-08-01', '2023-12-15', '2025-07-20']),
    'fecha_fin': pd.to_datetime(['2024-08-10', '2021-12-25', '2025-08-01'])
})
    df['diferencia_dias'] = df['fecha_fin'] - df['fecha_inicio']
    document.add_paragraph(df.to_string())
    
    document.add_paragraph("Con el pd.DateOffset() se puede  sumar o restar períodos de tiempo a fechas en un DataFrame o Series de tipo datetime:")
    df['fecha_inicio_mas_6d'] = df['fecha_inicio'] + pd.DateOffset(days=6)
    df['fecha_menos_1_mes'] = df['fecha_fin'] - pd.DateOffset(months=1)
    df['fecha_mas_2_años'] = df['fecha_inicio'] + pd.DateOffset(years=2)
    document.add_paragraph(df.to_string())

    fechas = pd.date_range(start='2024-01-01', end='2024-12-31', freq='D')
    datos = pd.Series(range(len(fechas)), index=fechas)

    document.add_paragraph("Con el metodo resample puedes reagrupar las fecha e integrarle una funcion de agregacion, ejemplo reagrupacion a final de mes resample('M') con promedio de dias ")
    datos_mensuales = datos.resample('M').mean()
    document.add_paragraph(datos_mensuales.to_string())

    document.add_paragraph("Para realizar promedios moviles se utiliza el metodo rolling(window=7).mean() donde el parametro window especifica de cuanto sera la ventana deslizante de 7 días sobre la serie de datos.")
    datos_rolling = datos.rolling(window=7).mean()
    document.add_paragraph(datos_rolling['2024-01-01':'2024-01-14'].to_string())

